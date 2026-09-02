#!/usr/bin/env python3
"""
Word document authoring with real tracked changes.

Two operations:

  create  — build a formatted .docx from Markdown-ish text
  revise  — write the NEXT version of an existing document, where every
            difference is a genuine Word revision (Accept/Reject in the
            Review ribbon), not coloured text pretending to be one.

Why the XML is written by hand: no Python library emits OOXML revision
markup. python-docx has no concept of it, and pandoc only *reads* tracked
changes. So `revise` diffs the two documents and emits `w:ins` / `w:del`
elements itself.

Reading a revised document back is the subtle part. python-docx's
`paragraph.text` walks only the `w:r` children directly under `w:p`, so runs
nested inside `w:ins` are invisible to it and `w:delText` is not `w:t`. Its
idea of "the text" is therefore neither the before nor the after. paragraph_text()
below reads a specific side of the revision instead, which is what makes
v2 -> v3 chain correctly: v3 diffs against v2's *final* text.
"""
from __future__ import annotations

import argparse
import copy
import datetime as _dt
import difflib
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

BODY_FONT = "Calibri"
HEAD_FONT = "Calibri Light"
ACCENT = RGBColor(0x1F, 0x38, 0x64)


# ---------------------------------------------------------------- markdown

def parse_blocks(text: str):
    """Markdown-ish source -> [(kind, level, text)] blocks.

    Supports ATX headings, '-'/'*' bullets, '1.' numbered items, '>' quotes,
    '---' rules, and blank-line-separated paragraphs. Deliberately small: this
    formats documents, it is not a CommonMark implementation.
    """
    blocks = []
    para: list[str] = []

    def flush():
        if para:
            blocks.append(("p", 0, " ".join(para).strip()))
            para.clear()

    for raw in text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            flush()
            continue
        if re.fullmatch(r"(-\s*){3,}|(\*\s*){3,}|(_\s*){3,}", stripped):
            flush()
            blocks.append(("rule", 0, ""))
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            flush()
            blocks.append(("h", len(m.group(1)), m.group(2).strip()))
            continue
        m = re.match(r"^[-*+]\s+(.*)$", stripped)
        if m:
            flush()
            blocks.append(("ul", 0, m.group(1).strip()))
            continue
        m = re.match(r"^\d+[.)]\s+(.*)$", stripped)
        if m:
            flush()
            blocks.append(("ol", 0, m.group(1).strip()))
            continue
        if stripped.startswith(">"):
            flush()
            blocks.append(("quote", 0, stripped.lstrip("> ").strip()))
            continue
        para.append(stripped)

    flush()
    return blocks


_INLINE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*[^*]+?\*|_[^_]+?_|`[^`]+?`)", re.S)


def split_inline(text: str):
    """Split on inline markers -> [(text, bold, italic, mono)]."""
    out = []
    for piece in _INLINE.split(text):
        if not piece:
            continue
        if (piece.startswith("**") and piece.endswith("**")) or (
            piece.startswith("__") and piece.endswith("__")
        ):
            out.append((piece[2:-2], True, False, False))
        elif piece.startswith("`") and piece.endswith("`"):
            out.append((piece[1:-1], False, False, True))
        elif (piece.startswith("*") and piece.endswith("*")) or (
            piece.startswith("_") and piece.endswith("_")
        ):
            out.append((piece[1:-1], False, True, False))
        else:
            out.append((piece, False, False, False))
    return out or [(text, False, False, False)]


# ---------------------------------------------------------------- styling

STYLE_FOR = {"h": None, "p": "Body Text", "ul": "List Bullet",
             "ol": "List Number", "quote": "Intense Quote", "rule": "Body Text"}


def block_style(kind: str, level: int) -> str:
    if kind == "h":
        return f"Heading {min(level, 4)}"
    return STYLE_FOR.get(kind) or "Body Text"


def apply_base_styles(doc: Document) -> None:
    """House style: readable body, restrained headings, sane margins."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before in (
        ("Heading 1", 18, 20), ("Heading 2", 14, 16),
        ("Heading 3", 12, 12), ("Heading 4", 11, 10),
    ):
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = HEAD_FONT
        st.font.size = Pt(size)
        st.font.bold = name in ("Heading 1", "Heading 2")
        st.font.color.rgb = ACCENT
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)
        section.top_margin = section.bottom_margin = Inches(1.0)


def add_page_numbers(doc: Document) -> None:
    """Footer with a PAGE field. Fields need raw XML — python-docx has none."""
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for kind, instr in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), kind)
            run._r.append(fld)
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = f" {instr} "
            run._r.append(it)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def write_blocks(doc: Document, blocks) -> None:
    for kind, level, text in blocks:
        if kind == "rule":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "BFBFBF")
            bdr.append(bottom)
            pPr.append(bdr)
            continue
        p = doc.add_paragraph(style=block_style(kind, level))
        for chunk, bold, italic, mono in split_inline(text):
            r = p.add_run(chunk)
            r.bold, r.italic = bold, italic
            if mono:
                r.font.name = "Consolas"
                r.font.size = Pt(10)


def create(md_text: str, out_path: str, title: str,
           subtitle: str = "", author: str = "Noto") -> str:
    doc = Document()
    apply_base_styles(doc)

    if title:
        t = doc.add_paragraph(title, style="Title")
        t.paragraph_format.space_after = Pt(4)
    if subtitle:
        s = doc.add_paragraph(subtitle, style="Subtitle")
        s.paragraph_format.space_after = Pt(2)
    if title:
        stamp = doc.add_paragraph(
            f"{author} · {_dt.date.today():%B %-d, %Y}")
        stamp.runs[0].font.size = Pt(9)
        stamp.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        stamp.paragraph_format.space_after = Pt(18)

    write_blocks(doc, parse_blocks(md_text))
    add_page_numbers(doc)

    doc.core_properties.title = title or os.path.basename(out_path)
    doc.core_properties.author = author

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return out_path


# ------------------------------------------------------- revision plumbing

def paragraph_text(p, side: str = "final") -> str:
    """Text of one paragraph on a given side of its revisions.

    side='final'    -> as if every change were accepted (insertions kept)
    side='original' -> as if every change were rejected (deletions restored)

    Needed because python-docx's own .text sees neither: it reads only the
    w:r children directly under w:p, skipping anything inside w:ins, and
    deleted text lives in w:delText rather than w:t.
    """
    out = []
    for node in p.iter():
        tag = node.tag
        if tag == W + "t":
            parent = node.getparent().getparent()
            in_ins = parent is not None and parent.tag == W + "ins"
            if side == "original" and in_ins:
                continue
            out.append(node.text or "")
        elif tag == W + "delText":
            if side == "final":
                continue
            out.append(node.text or "")
    return "".join(out)


BYLINE_RE = re.compile(r"^.+ · \w+ \d{1,2}, \d{4}( · revised)?$")


def is_front_matter(style: str, text: str) -> bool:
    """Title, subtitle and the generated byline are chrome, not content.

    They are re-emitted on every revision, so including them in the diff
    would delete and re-insert the whole heading block each time.
    """
    return style in ("Title", "Subtitle") or bool(BYLINE_RE.match(text))


def doc_blocks(path: str, side: str = "final"):
    """Read a .docx back into (style_name, text) pairs, revisions resolved."""
    doc = Document(path)
    items = []
    for p in doc.paragraphs:
        txt = paragraph_text(p._p, side).strip()
        if not txt:
            continue
        style = p.style.name if p.style is not None else "Body Text"
        if is_front_matter(style, txt):
            continue
        items.append((style, txt))
    return items


class _Rev:
    """Allocates the w:id values Word requires to be unique per revision."""

    def __init__(self, author: str, when: str | None = None):
        self.author = author
        self.date = when or _dt.datetime.now(
            _dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
        self._n = 1000

    def next_id(self) -> str:
        self._n += 1
        return str(self._n)

    def _stamp(self, el):
        el.set(qn("w:id"), self.next_id())
        el.set(qn("w:author"), self.author)
        el.set(qn("w:date"), self.date)
        return el

    def ins(self, runs):
        el = self._stamp(OxmlElement("w:ins"))
        for r in runs:
            el.append(r)
        return el

    def dele(self, runs):
        el = self._stamp(OxmlElement("w:del"))
        for r in runs:
            el.append(r)
        return el

    def mark_para(self, p, tag: str):
        """Mark the paragraph MARK itself inserted/deleted.

        Without this Word shows the text change but not the paragraph break,
        so accepting the revision leaves a stray empty paragraph behind.
        """
        pPr = p.get_or_add_pPr()
        rPr = pPr.find(qn("w:rPr"))
        if rPr is None:
            # CT_PPr is a sequence: w:pStyle first, w:rPr LAST. Inserting at
            # index 0 puts rPr ahead of pStyle, which is schema-invalid and
            # makes Word report the file as corrupt. Append instead.
            rPr = OxmlElement("w:rPr")
            pPr.append(rPr)
        rPr.append(self._stamp(OxmlElement(f"w:{tag}")))


def _run(text: str, deleted: bool = False):
    r = OxmlElement("w:r")
    t = OxmlElement("w:delText" if deleted else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r


_TOKEN = re.compile(r"\S+\s*")


def _tokens(s: str):
    return _TOKEN.findall(s) or ([s] if s else [])


def _diff_into(p, old: str, new: str, rev: _Rev) -> None:
    """Word-level intra-paragraph diff, appended to p as ins/del runs."""
    a, b = _tokens(old), _tokens(new)
    for op, i1, i2, j1, j2 in difflib.SequenceMatcher(None, a, b).get_opcodes():
        if op == "equal":
            p.append(_run("".join(a[i1:i2])))
        elif op == "delete":
            p.append(rev.dele([_run("".join(a[i1:i2]), deleted=True)]))
        elif op == "insert":
            p.append(rev.ins([_run("".join(b[j1:j2]))]))
        else:  # replace -> delete then insert, the order Word expects
            p.append(rev.dele([_run("".join(a[i1:i2]), deleted=True)]))
            p.append(rev.ins([_run("".join(b[j1:j2]))]))


def revise(base_path: str, new_md: str, out_path: str,
           author: str = "Noto", title: str | None = None) -> dict:
    """Write out_path as base_path + tracked changes reaching new_md."""
    old = doc_blocks(base_path, side="final")
    new = [(block_style(k, lv), t) for k, lv, t in parse_blocks(new_md)
           if k != "rule"]

    doc = Document()
    apply_base_styles(doc)
    rev = _Rev(author)

    if title:
        doc.add_paragraph(title, style="Title")
        stamp = doc.add_paragraph(
            f"{author} · {_dt.date.today():%B %-d, %Y} · revised")
        stamp.runs[0].font.size = Pt(9)
        stamp.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        stamp.paragraph_format.space_after = Pt(18)

    counts = {"unchanged": 0, "inserted": 0, "deleted": 0, "edited": 0}
    sm = difflib.SequenceMatcher(None, [t for _, t in old], [t for _, t in new])

    for op, i1, i2, j1, j2 in sm.get_opcodes():
        if op == "equal":
            for style, text in new[j1:j2]:
                para = doc.add_paragraph(style=style)
                para._p.append(_run(text))
                counts["unchanged"] += 1
        elif op == "insert":
            for style, text in new[j1:j2]:
                para = doc.add_paragraph(style=style)
                para._p.append(rev.ins([_run(text)]))
                rev.mark_para(para._p, "ins")
                counts["inserted"] += 1
        elif op == "delete":
            for style, text in old[i1:i2]:
                para = doc.add_paragraph(style=style)
                para._p.append(rev.dele([_run(text, deleted=True)]))
                rev.mark_para(para._p, "del")
                counts["deleted"] += 1
        else:
            # Pair them up so a reworded paragraph reads as an edit rather
            # than a wholesale delete + insert.
            olds, news = old[i1:i2], new[j1:j2]
            for k in range(max(len(olds), len(news))):
                if k < len(olds) and k < len(news):
                    style = news[k][0]
                    para = doc.add_paragraph(style=style)
                    _diff_into(para._p, olds[k][1], news[k][1], rev)
                    counts["edited"] += 1
                elif k < len(news):
                    para = doc.add_paragraph(style=news[k][0])
                    para._p.append(rev.ins([_run(news[k][1])]))
                    rev.mark_para(para._p, "ins")
                    counts["inserted"] += 1
                else:
                    para = doc.add_paragraph(style=olds[k][0])
                    para._p.append(rev.dele([_run(olds[k][1], deleted=True)]))
                    rev.mark_para(para._p, "del")
                    counts["deleted"] += 1

    add_page_numbers(doc)
    doc.core_properties.author = author
    doc.core_properties.title = title or os.path.basename(out_path)
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return counts


# ---------------------------------------------------------------- versions

VERSION_RE = re.compile(r"^(?P<stem>.+?)[ _-]+v(?P<n>\d+)$", re.I)


def parse_version(filename: str):
    stem = os.path.splitext(os.path.basename(filename))[0]
    m = VERSION_RE.match(stem)
    return (m.group("stem").strip(), int(m.group("n"))) if m else (stem, 1)


def version_path(folder: str, stem: str, n: int) -> str:
    return os.path.join(folder, f"{stem} v{n}.docx")


def latest_version(folder: str, stem: str):
    """Highest existing version of `stem` in folder, or None."""
    if not os.path.isdir(folder):
        return None
    best = None
    for name in os.listdir(folder):
        if not name.lower().endswith(".docx") or name.startswith("~$"):
            continue
        s, n = parse_version(name)
        if s.lower() == stem.lower() and (best is None or n > best[1]):
            best = (os.path.join(folder, name), n)
    return best


# -------------------------------------------------------------------- cli

def main() -> int:
    ap = argparse.ArgumentParser(description="Create and revise Word documents")
    sub = ap.add_subparsers(dest="cmd", required=True)

    c = sub.add_parser("create", help="create v1 of a document")
    c.add_argument("--folder", required=True)
    c.add_argument("--name", required=True, help="document stem, no version")
    c.add_argument("--title")
    c.add_argument("--subtitle", default="")
    c.add_argument("--author", default="Noto")
    c.add_argument("--input", help="Markdown file (default: stdin)")

    r = sub.add_parser("revise", help="write the next version, tracked")
    r.add_argument("--folder", required=True)
    r.add_argument("--name", required=True)
    r.add_argument("--author", default="Noto")
    r.add_argument("--title")
    r.add_argument("--input", help="Markdown file (default: stdin)")

    s = sub.add_parser("versions", help="list versions of a document")
    s.add_argument("--folder", required=True)
    s.add_argument("--name", required=True)

    a = ap.parse_args()

    if a.cmd == "versions":
        found = latest_version(a.folder, a.name)
        if not found:
            print(f"no versions of '{a.name}' in {a.folder}")
            return 1
        print(f"latest: {found[0]} (v{found[1]})")
        return 0

    md = open(a.input).read() if a.input else sys.stdin.read()

    if a.cmd == "create":
        existing = latest_version(a.folder, a.name)
        if existing:
            print(f"ERROR: {a.name} already exists at v{existing[1]} "
                  f"({existing[0]}).\n"
                  f"Use 'revise' so the change is tracked, not 'create'.",
                  file=sys.stderr)
            return 1
        out = version_path(a.folder, a.name, 1)
        create(md, out, a.title or a.name, a.subtitle, a.author)
        print(out)
        return 0

    found = latest_version(a.folder, a.name)
    if not found:
        print(f"ERROR: no existing version of '{a.name}' in {a.folder}. "
              f"Use 'create' first.", file=sys.stderr)
        return 1
    base, n = found
    out = version_path(a.folder, a.name, n + 1)
    counts = revise(base, md, out, a.author, a.title or a.name)
    print(out)
    print(f"v{n} -> v{n+1}: {counts['inserted']} inserted, "
          f"{counts['deleted']} deleted, {counts['edited']} edited, "
          f"{counts['unchanged']} unchanged", file=sys.stderr)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
